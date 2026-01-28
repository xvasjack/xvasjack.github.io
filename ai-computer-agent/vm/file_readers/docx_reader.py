"""
DOCX Reader - Analyze Word documents for DD report quality checking.

This module extracts:
- Cover page fields (title, company, preparedFor, purpose, confidential)
- Section headers and numbering
- Tables (headers, rows, column counts)
- Images/figures
- Text content for analysis
"""

import os
import re
from typing import List, Optional, Dict, Any
from dataclasses import dataclass, field
import logging

try:
    from docx import Document
    from docx.shared import Inches, Pt
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("Missing python-docx. Run: pip install python-docx")

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("docx_reader")


@dataclass
class TableInfo:
    """Information about a table in the document"""
    index: int
    location: str  # e.g., "After section 2.1"
    header_row: List[str]
    row_count: int
    column_count: int
    sample_rows: List[List[str]] = field(default_factory=list)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "index": self.index,
            "location": self.location,
            "header_row": self.header_row,
            "row_count": self.row_count,
            "column_count": self.column_count,
            "sample_rows": self.sample_rows[:3],
        }


@dataclass
class SectionInfo:
    """Information about a section header"""
    number: str  # e.g., "4.0", "4.9", "8"
    title: str
    level: int  # 1=heading1, 2=heading2, 3=heading3
    full_text: str  # Complete header text

    def to_dict(self) -> Dict[str, Any]:
        return {
            "number": self.number,
            "title": self.title,
            "level": self.level,
            "full_text": self.full_text,
        }


@dataclass
class CoverPageInfo:
    """Cover page information for DD reports"""
    title: Optional[str] = None
    company_name: Optional[str] = None
    prepared_for: Optional[str] = None
    purpose: Optional[str] = None
    confidential: Optional[str] = None
    date: Optional[str] = None

    def to_dict(self) -> Dict[str, Any]:
        return {
            "title": self.title,
            "company_name": self.company_name,
            "prepared_for": self.prepared_for,
            "purpose": self.purpose,
            "confidential": self.confidential,
            "date": self.date,
        }


@dataclass
class DOCXAnalysis:
    """Complete analysis of a DOCX file"""
    file_path: str
    paragraph_count: int = 0
    table_count: int = 0
    image_count: int = 0
    cover_page: CoverPageInfo = field(default_factory=CoverPageInfo)
    sections: List[SectionInfo] = field(default_factory=list)
    tables: List[TableInfo] = field(default_factory=list)
    all_text: List[str] = field(default_factory=list)
    issues: List[str] = field(default_factory=list)
    summary: str = ""

    def to_dict(self) -> Dict[str, Any]:
        return {
            "file_path": self.file_path,
            "paragraph_count": self.paragraph_count,
            "table_count": self.table_count,
            "image_count": self.image_count,
            "cover_page": self.cover_page.to_dict(),
            "sections": [s.to_dict() for s in self.sections],
            "tables": [t.to_dict() for t in self.tables],
            "issues": self.issues,
            "summary": self.summary,
        }


def analyze_docx(file_path: str) -> DOCXAnalysis:
    """
    Analyze a DOCX file and extract DD report structure.

    Args:
        file_path: Path to the DOCX file

    Returns:
        DOCXAnalysis with extracted information
    """
    logger.info(f"Analyzing DOCX: {file_path}")

    if not HAS_DOCX:
        return DOCXAnalysis(
            file_path=file_path,
            issues=["python-docx not installed"],
            summary="Error: python-docx not installed"
        )

    if not os.path.exists(file_path):
        return DOCXAnalysis(
            file_path=file_path,
            issues=["File not found"],
            summary="Error: File not found"
        )

    try:
        doc = Document(file_path)
    except Exception as e:
        return DOCXAnalysis(
            file_path=file_path,
            issues=[f"Could not open file: {e}"],
            summary=f"Error opening file: {e}"
        )

    # Category 2 fix: Wrap in try/finally to ensure document cleanup
    try:
        analysis = DOCXAnalysis(file_path=file_path)

        # Extract all text from paragraphs
        all_paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                all_paragraphs.append(text)
                analysis.all_text.append(text)

        analysis.paragraph_count = len(all_paragraphs)

        # Extract cover page info (first few paragraphs before section 1.0)
        analysis.cover_page = extract_cover_page(all_paragraphs)

        # Extract sections
        analysis.sections = extract_sections(doc)

        # Extract tables
        analysis.tables = extract_tables(doc, analysis.sections)
        analysis.table_count = len(analysis.tables)

        # Count images
        analysis.image_count = count_images(doc)

        # Check for issues
        analysis.issues = check_quality_issues(analysis)

        # Generate summary
        analysis.summary = generate_summary(analysis)

        return analysis
    finally:
        # Category 2 fix: Document cleanup - python-docx doesn't have explicit close
        # but we ensure any internal resources are released
        del doc


def extract_cover_page(paragraphs: List[str]) -> CoverPageInfo:
    """Extract cover page information from the first paragraphs"""
    cover = CoverPageInfo()

    # Look for cover page elements in the first ~10 paragraphs
    cover_texts = paragraphs[:15] if len(paragraphs) > 15 else paragraphs

    for i, text in enumerate(cover_texts):
        text_lower = text.lower()

        # Check for title (PRE-DUE DILIGENCE REPORT or similar)
        if "due diligence" in text_lower and "report" in text_lower:
            cover.title = text

        # Check for "Prepared for"
        elif "prepared for" in text_lower:
            cover.prepared_for = text

        # Check for purpose (contains "evaluation" or "investment")
        elif "evaluation" in text_lower or "investment" in text_lower:
            cover.purpose = text

        # Check for confidential disclaimer
        elif "confidential" in text_lower:
            cover.confidential = text

        # Company name is usually between title and "Prepared for"
        # Look for company name patterns
        # Category 12 fix: Company name extraction - be more permissive
        elif cover.title and not cover.prepared_for:
            # Likely company name if it's after title but before prepared for
            company_suffixes = ["pte", "ltd", "inc", "corp", "llc", "gmbh", "co.", "company", "limited"]
            if any(suffix in text_lower for suffix in company_suffixes):
                cover.company_name = text
            elif len(text) > 3 and not text.startswith("1.") and not text.startswith("2."):
                # Could be company name if short and not a section header
                if not cover.company_name and not any(kw in text_lower for kw in ["confidential", "prepared", "purpose", "evaluation"]):
                    cover.company_name = text

    return cover


def extract_sections(doc) -> List[SectionInfo]:
    """Extract section headers and their hierarchy"""
    sections = []

    # Pattern for section numbers like "1.0", "1.1", "4.9", "8", "8."
    section_pattern = re.compile(r'^(\d+(?:\.\d+)?)\s*[.\s]*(.*)$')

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Check if this is a heading style
        # Category 12 fix: Heading style detection should be case-insensitive
        style_name = para.style.name if para.style else ""
        style_lower = style_name.lower()
        is_heading = "heading" in style_lower

        # Also check for section number pattern
        match = section_pattern.match(text)

        if match or is_heading:
            if match:
                number = match.group(1)
                title = match.group(2).strip()
            else:
                # Try to extract number from text
                number = ""
                title = text

            # Determine level
            level = 1
            if "Heading 2" in style_name or (number and "." in number and number.count(".") == 1):
                level = 2
            elif "Heading 3" in style_name or (number and number.count(".") >= 2):
                level = 3

            sections.append(SectionInfo(
                number=number,
                title=title,
                level=level,
                full_text=text
            ))

    return sections


def extract_tables(doc, sections: List[SectionInfo]) -> List[TableInfo]:
    """Extract table information"""
    tables = []

    # Map paragraph positions to sections
    section_map = {}
    current_section = "Before content"

    for table_idx, table in enumerate(doc.tables):
        # Get header row
        header_row = []
        num_rows = len(table.rows)
        if num_rows > 0:
            for cell in table.rows[0].cells:
                header_row.append(cell.text.strip())

        # Get sample rows - safely slice within bounds
        sample_rows = []
        # Category 4 fix: Explicit bounds check for table row slicing
        end_row = min(4, num_rows)  # Don't exceed actual row count
        for row_idx in range(1, end_row):  # Start from row 1 (skip header)
            row = table.rows[row_idx]
            row_data = [cell.text.strip() for cell in row.cells]
            sample_rows.append(row_data)

        # Try to determine location based on nearby text
        location = f"Table {table_idx + 1}"

        tables.append(TableInfo(
            index=table_idx,
            location=location,
            header_row=header_row,
            row_count=len(table.rows) - 1,  # Exclude header
            column_count=len(header_row),
            sample_rows=sample_rows
        ))

    return tables


def count_images(doc) -> int:
    """Count images in the document"""
    image_count = 0

    # Check inline shapes
    for para in doc.paragraphs:
        for run in para.runs:
            # Category 5 fix: Image counting errors should be logged, not silently swallowed
            try:
                if run._element.xpath('.//a:blip'):
                    image_count += 1
            except Exception as e:
                logger.debug(f"Error checking image in run: {e}")

    # Check shapes in document
    try:
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1
    except Exception as e:
        logger.debug(f"Error counting document images: {e}")

    return image_count


def check_quality_issues(analysis: DOCXAnalysis) -> List[str]:
    """Check for quality issues in the DD report"""
    issues = []

    # Cover page checks
    if not analysis.cover_page.title:
        issues.append("Missing cover page title")
    elif "PRE-DUE DILIGENCE" not in analysis.cover_page.title.upper():
        issues.append(f"Cover page title should be 'PRE-DUE DILIGENCE REPORT', got: '{analysis.cover_page.title}'")

    if not analysis.cover_page.prepared_for:
        issues.append("Missing 'Prepared for' line on cover page")

    if not analysis.cover_page.purpose:
        issues.append("Missing purpose statement on cover page")

    if not analysis.cover_page.confidential:
        issues.append("Missing confidential disclaimer on cover page")

    # Section numbering checks
    section_numbers = {s.number for s in analysis.sections if s.number}

    # Check for correct financial section number (should be 4.0 not 3.0)
    has_financials = any("financial" in s.title.lower() for s in analysis.sections)
    if has_financials:
        financials_section = next((s for s in analysis.sections if "financial" in s.title.lower()), None)
        if financials_section and financials_section.number:
            if not financials_section.number.startswith("4"):
                issues.append(f"Financials should be section 4.0, got: {financials_section.number}")

    # Check for Pre-DD Workplan section number (should be 4.9)
    workplan_section = next((s for s in analysis.sections if "workplan" in s.title.lower() or "pre-dd" in s.title.lower()), None)
    if workplan_section and workplan_section.number:
        if workplan_section.number != "4.9":
            issues.append(f"Pre-DD Workplan should be section 4.9, got: {workplan_section.number}")

    # Check for Future Plans section number (should be 8)
    # Category 6 fix: "Future" section detection too broad - be more specific
    future_section = next(
        (s for s in analysis.sections if "future" in s.title.lower() and ("plan" in s.title.lower() or "growth" in s.title.lower() or "outlook" in s.title.lower())),
        None
    )
    if future_section and future_section.number:
        if not future_section.number.startswith("8"):
            issues.append(f"Future Plans should be section 8, got: {future_section.number}")

    # Table checks
    for table in analysis.tables:
        # Check for SGD ('000) format in financial tables
        headers_lower = [h.lower() for h in table.header_row]
        if any("sgd" in h or "revenue" in h or "profit" in h for h in headers_lower):
            has_thousands_format = any("'000" in h or "000)" in h for h in table.header_row)
            if not has_thousands_format:
                issues.append(f"Financial table {table.index + 1} should use 'SGD ('000)' format")

    # Check for Competition Landscape table
    competition_table = next((t for t in analysis.tables if any("segment" in h.lower() for h in t.header_row)), None)
    if competition_table:
        # Check for CAGR in headers
        has_cagr = any("cagr" in h.lower() for h in competition_table.header_row)
        if not has_cagr:
            issues.append("Competition Landscape table should have 'Growth (CAGR %)' column")

    return issues


def generate_summary(analysis: DOCXAnalysis) -> str:
    """Generate a summary of the analysis"""
    summary = f"""DOCX Analysis Summary:
- File: {os.path.basename(analysis.file_path)}
- Paragraphs: {analysis.paragraph_count}
- Tables: {analysis.table_count}
- Images: {analysis.image_count}
- Sections found: {len(analysis.sections)}
- Issues found: {len(analysis.issues)}

Cover Page:
- Title: {analysis.cover_page.title or 'MISSING'}
- Company: {analysis.cover_page.company_name or 'MISSING'}
- Prepared for: {analysis.cover_page.prepared_for or 'MISSING'}
- Purpose: {'Present' if analysis.cover_page.purpose else 'MISSING'}
- Confidential: {'Present' if analysis.cover_page.confidential else 'MISSING'}

Sections:
{chr(10).join(f'- {s.number} {s.title}' for s in analysis.sections[:10])}
{'...' if len(analysis.sections) > 10 else ''}

Issues:
{chr(10).join('- ' + issue for issue in analysis.issues) if analysis.issues else '- None'}
"""
    return summary


def find_table_by_header(analysis: DOCXAnalysis, header_keyword: str) -> Optional[TableInfo]:
    """Find a table by a keyword in its header row"""
    for table in analysis.tables:
        if any(header_keyword.lower() in h.lower() for h in table.header_row):
            return table
    return None


def get_section_by_title(analysis: DOCXAnalysis, title_keyword: str) -> Optional[SectionInfo]:
    """Find a section by a keyword in its title"""
    for section in analysis.sections:
        if title_keyword.lower() in section.title.lower():
            return section
    return None


# =============================================================================
# DD REPORT SPECIFIC CHECKS
# =============================================================================


def check_dd_report_compliance(analysis: DOCXAnalysis) -> Dict[str, Any]:
    """
    Check DD report against template requirements.
    Returns detailed compliance results.
    """
    results = {
        "passed": True,
        "cover_page": {},
        "sections": {},
        "tables": {},
        "predd_workplan": {},
        "issues": [],
    }

    # Cover page checks
    results["cover_page"] = {
        "has_title": analysis.cover_page.title is not None,
        "title_correct": analysis.cover_page.title and "PRE-DUE DILIGENCE" in analysis.cover_page.title.upper(),
        "has_company": analysis.cover_page.company_name is not None,
        "has_prepared_for": analysis.cover_page.prepared_for is not None,
        "has_purpose": analysis.cover_page.purpose is not None,
        "has_confidential": analysis.cover_page.confidential is not None,
    }

    # Check section numbers
    financials = get_section_by_title(analysis, "Financial")
    workplan = get_section_by_title(analysis, "Workplan") or get_section_by_title(analysis, "Pre-DD")
    future = get_section_by_title(analysis, "Future")

    results["sections"] = {
        "financials_number": financials.number if financials else None,
        "financials_correct": financials and financials.number.startswith("4"),
        "workplan_number": workplan.number if workplan else None,
        "workplan_correct": workplan and workplan.number == "4.9",
        "future_number": future.number if future else None,
        "future_correct": future and future.number.startswith("8"),
    }

    # Check Competition Landscape table
    competition_table = find_table_by_header(analysis, "Segment")
    if competition_table:
        results["tables"]["competition_landscape"] = {
            "found": True,
            "column_count": competition_table.column_count,
            "has_cagr": any("cagr" in h.lower() for h in competition_table.header_row),
            "headers": competition_table.header_row,
        }
    else:
        results["tables"]["competition_landscape"] = {"found": False}

    # Check financial tables
    financial_tables = [t for t in analysis.tables if any("sgd" in h.lower() for h in t.header_row)]
    results["tables"]["financial_tables"] = {
        "count": len(financial_tables),
        "use_thousands_format": all(
            any("'000" in h or "000)" in h for h in t.header_row)
            for t in financial_tables
        ) if financial_tables else False,
    }

    # Check Pre-DD Workplan
    workplan_table = find_table_by_header(analysis, "Consideration") or find_table_by_header(analysis, "Evidence")
    if workplan_table:
        expected_rows = [
            "customer analysis",
            "pipeline analysis",
            "pricing power",
            "unit economics",
            "billing",
            "forecast",
            "partner ecosystem",
        ]
        found_rows = []
        for row in workplan_table.sample_rows:
            row_text = " ".join(row).lower()
            for expected in expected_rows:
                if expected in row_text:
                    found_rows.append(expected)

        results["predd_workplan"] = {
            "found": True,
            "row_count": workplan_table.row_count,
            "standard_rows_found": len(set(found_rows)),
            "expected_rows": 7,
        }
    else:
        results["predd_workplan"] = {"found": False}

    # Determine overall pass/fail
    critical_checks = [
        results["cover_page"]["title_correct"],
        results["sections"]["financials_correct"],
        results["sections"]["workplan_correct"],
        results["sections"]["future_correct"],
    ]
    non_none = [c for c in critical_checks if c is not None]
    results["passed"] = bool(non_none) and all(non_none)
    results["issues"] = analysis.issues

    return results


# =============================================================================
# ENTRY POINT
# =============================================================================


if __name__ == "__main__":
    import sys
    import json

    if len(sys.argv) < 2:
        print("Usage: python docx_reader.py <file.docx>")
        sys.exit(1)

    analysis = analyze_docx(sys.argv[1])
    print(analysis.summary)
    print("\n--- Compliance Check ---")
    compliance = check_dd_report_compliance(analysis)
    print(json.dumps(compliance, indent=2))
